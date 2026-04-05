import asyncio
import io
import json
import logging
import os
import re
import socket
import sys
import time
import threading
from dataclasses import dataclass
from ipaddress import ip_address
from typing import Any


def _load_app_settings() -> None:
    path = os.getenv("APP_CONFIG_PATH")
    if not path:
        path = os.path.join(os.path.dirname(__file__), "appsettings.json")
    try:
        if not os.path.isfile(path):
            return
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if not isinstance(data, dict):
            return
        for key, value in data.items():
            if value is None:
                continue
            if os.getenv(str(key)) not in (None, ""):
                continue
            os.environ[str(key)] = str(value)

        app_root = os.getenv("APP_ROOT")
        if app_root and os.path.isdir(app_root):
            try:
                if os.path.abspath(os.getcwd()) != os.path.abspath(app_root):
                    os.chdir(app_root)
                if app_root not in sys.path:
                    sys.path.insert(0, app_root)
                if not os.getenv("PYTHONPATH"):
                    os.environ["PYTHONPATH"] = app_root
            except Exception as exc:
                print(f"[config] Failed to apply APP_ROOT {app_root}: {exc}", flush=True)
    except Exception as exc:
        print(f"[config] Failed to load {path}: {exc}", flush=True)


_load_app_settings()


import httpx
import numpy as np
import pdfplumber
from fastapi import FastAPI, File, Header, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import PlainTextResponse, StreamingResponse
from pydantic import BaseModel, HttpUrl
from PIL import Image

from module import LayoutRecognizer, OCR, TableStructureRecognizer

try:
    from docx import Document
except Exception:
    Document = None


logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO").upper(),
    format="%(asctime)s | %(levelname)s | %(message)s",
)


@dataclass(frozen=True)
class Settings:
    api_key: str = os.getenv("API_KEY", "")
    max_upload_mb: int = int(os.getenv("MAX_UPLOAD_MB", "25"))
    max_pages: int = int(os.getenv("MAX_PAGES", "50"))
    max_concurrent_jobs: int = int(os.getenv("MAX_CONCURRENT_JOBS", "1"))
    default_threshold: float = float(os.getenv("DEFAULT_THRESHOLD", "0.2"))
    cors_allow_origins: str = os.getenv("CORS_ALLOW_ORIGINS", "")
    # SSRF guard for URL-based extraction. Set to 0 only if you know what you are doing.
    block_private_networks: bool = os.getenv("BLOCK_PRIVATE_NETWORKS", "1").strip().lower() in {"1", "true", "yes"}


SETTINGS = Settings()
JOB_SEM = asyncio.Semaphore(max(1, SETTINGS.max_concurrent_jobs))
PDF_LOCK = threading.Lock()


app = FastAPI(title="DeepDoc+VietOCR API", version="1.0.0")

if SETTINGS.cors_allow_origins:
    app.add_middleware(
        CORSMiddleware,
        allow_origins=[o.strip() for o in SETTINGS.cors_allow_origins.split(",") if o.strip()],
        allow_credentials=False,
        allow_methods=["*"],
        allow_headers=["*"],
    )


@app.get("/health")
def health() -> dict[str, Any]:
    return {"ok": True}


_ocr: OCR | None = None
_layout: Any | None = None
_tsr: Any | None = None


def _get_ocr() -> OCR:
    global _ocr
    if _ocr is None:
        # OCR uses onnx/det.onnx + VietOCR weights in vietocr_assets/weight/...
        _ocr = OCR()
    return _ocr


def _get_layout():
    global _layout
    if _layout is None:
        _layout = LayoutRecognizer("layout")
    return _layout


def _get_tsr():
    global _tsr
    if _tsr is None:
        _tsr = TableStructureRecognizer()
    return _tsr


def _require_api_key(x_api_key: str | None) -> None:
    if not SETTINGS.api_key:
        return
    if not x_api_key or x_api_key != SETTINGS.api_key:
        raise HTTPException(status_code=401, detail="Invalid API key")


def _check_upload_size(raw: bytes) -> None:
    if len(raw) > SETTINGS.max_upload_mb * 1024 * 1024:
        raise HTTPException(status_code=413, detail=f"File too large (>{SETTINGS.max_upload_mb}MB)")


def _is_private_host(host: str) -> bool:
    if not SETTINGS.block_private_networks:
        return False
    if not host:
        return True
    if host.lower() in {"localhost", "127.0.0.1", "::1"}:
        return True
    try:
        for family, _, _, _, sockaddr in socket.getaddrinfo(host, None):
            if family == socket.AF_INET:
                ip = ip_address(sockaddr[0])
            elif family == socket.AF_INET6:
                ip = ip_address(sockaddr[0])
            else:
                continue
            if (
                ip.is_private
                or ip.is_loopback
                or ip.is_link_local
                or ip.is_multicast
                or ip.is_reserved
                or ip.is_unspecified
            ):
                return True
    except Exception:
        return True
    return False


def _download_url(url: str) -> tuple[str, bytes]:
    parsed = httpx.URL(url)
    if parsed.scheme not in {"http", "https"}:
        raise HTTPException(status_code=400, detail="Only http/https URLs are supported")
    if _is_private_host(parsed.host or ""):
        raise HTTPException(status_code=400, detail="Blocked URL host (private/loopback networks are not allowed)")

    max_bytes = SETTINGS.max_upload_mb * 1024 * 1024
    headers = {"User-Agent": "deepdoc-vietocr-api/1.0"}

    try:
        with httpx.Client(follow_redirects=True, timeout=30.0, headers=headers) as client:
            with client.stream("GET", url) as resp:
                resp.raise_for_status()
                chunks: list[bytes] = []
                total = 0
                for chunk in resp.iter_bytes():
                    if not chunk:
                        continue
                    total += len(chunk)
                    if total > max_bytes:
                        raise HTTPException(
                            status_code=413, detail=f"Remote file too large (>{SETTINGS.max_upload_mb}MB)"
                        )
                    chunks.append(chunk)
                raw = b"".join(chunks)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to download URL: {e}")

    filename = os.path.basename(parsed.path or "") or "download"
    # If URL has no extension, infer from magic bytes (PDF) to help routing.
    if "." not in filename:
        if raw.startswith(b"%PDF"):
            filename += ".pdf"
        else:
            filename += ".jpg"
    return filename, raw


def _load_images(filename: str, raw: bytes) -> list[Image.Image]:
    ext = (os.path.splitext(filename)[1] or "").lower().lstrip(".")
    if ext == "pdf":
        # pdfplumber rendering isn't perfectly thread-safe; serialize to avoid crashes.
        with PDF_LOCK:
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                pages = []
                max_pages = min(len(pdf.pages), SETTINGS.max_pages)
                for i in range(max_pages):
                    page = pdf.pages[i]
                    pages.append(page.to_image(resolution=72 * 3).annotated.convert("RGB"))
                return pages

    try:
        img = Image.open(io.BytesIO(raw)).convert("RGB")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Unsupported/invalid image: {e}")
    return [img]


def _bbox_from_quad(quad: list[list[float]]) -> list[float]:
    xs = [p[0] for p in quad]
    ys = [p[1] for p in quad]
    return [float(min(xs)), float(min(ys)), float(max(xs)), float(max(ys))]


def _ocr_page(ocr: OCR, img: Image.Image) -> dict[str, Any]:
    arr = np.array(img)
    raw = ocr(arr, 0)
    lines = []
    texts = []
    for quad, rec in raw:
        text = rec[0]
        if not text:
            continue
        texts.append(text)
        lines.append(
            {
                "text": text,
                "quad": quad,
                "bbox": _bbox_from_quad(quad),
                "score": float(rec[1]),
            }
        )
    return {"text": "\n".join(texts), "lines": lines}


def _table_markdown(img: Image.Image, tb_components: list[dict[str, Any]], ocr_boxes: list[dict[str, Any]]) -> str:
    # This mirrors the logic in t_recognizer.py:get_table_markdown()
    boxes = LayoutRecognizer.sort_Y_firstly(
        ocr_boxes,
        np.mean([b["bottom"] - b["top"] for b in ocr_boxes]) / 3 if ocr_boxes else 0,
    )

    def gather(pattern: str, fzy: int = 10, ption: float = 0.6):
        eles = LayoutRecognizer.sort_Y_firstly([r for r in tb_components if re.match(pattern, r["label"])], fzy)
        eles = LayoutRecognizer.layouts_cleanup(boxes, eles, 5, ption)
        return LayoutRecognizer.sort_Y_firstly(eles, 0)

    headers = gather(r".*header$")
    rows = gather(r".* (row|header)")
    spans = gather(r".*spanning")
    clmns = sorted([r for r in tb_components if re.match(r"table column$", r["label"])], key=lambda x: x["x0"])
    clmns = LayoutRecognizer.layouts_cleanup(boxes, clmns, 5, 0.5)

    for b in boxes:
        ii = LayoutRecognizer.find_overlapped_with_threashold(b, rows, thr=0.3)
        if ii is not None:
            b["R"] = ii
            b["R_top"] = rows[ii]["top"]
            b["R_bott"] = rows[ii]["bottom"]

        ii = LayoutRecognizer.find_overlapped_with_threashold(b, headers, thr=0.3)
        if ii is not None:
            b["H_top"] = headers[ii]["top"]
            b["H_bott"] = headers[ii]["bottom"]
            b["H_left"] = headers[ii]["x0"]
            b["H_right"] = headers[ii]["x1"]
            b["H"] = ii

        ii = LayoutRecognizer.find_horizontally_tightest_fit(b, clmns)
        if ii is not None:
            b["C"] = ii
            b["C_left"] = clmns[ii]["x0"]
            b["C_right"] = clmns[ii]["x1"]

        ii = LayoutRecognizer.find_overlapped_with_threashold(b, spans, thr=0.3)
        if ii is not None:
            b["H_top"] = spans[ii]["top"]
            b["H_bott"] = spans[ii]["bottom"]
            b["H_left"] = spans[ii]["x0"]
            b["H_right"] = spans[ii]["x1"]
            b["SP"] = ii

    return TableStructureRecognizer.construct_table(boxes, markdown=True)


def _parse_markdown_table(md: str) -> list[list[str]] | None:
    if not md:
        return None
    lines = [l.strip() for l in md.strip().splitlines() if l.strip()]
    if len(lines) < 2:
        return None

    rows: list[list[str]] = []
    for line in lines:
        # Skip separator line like: | --- | --- |
        if re.match(r"^\|?(\s*:?-+:?\s*\|)+\s*$", line):
            continue
        if line.startswith("|"):
            line = line[1:]
        if line.endswith("|"):
            line = line[:-1]
        cells = [c.strip() for c in line.split("|")]
        rows.append(cells)

    if not rows:
        return None
    max_cols = max(len(r) for r in rows)
    for r in rows:
        if len(r) < max_cols:
            r.extend([""] * (max_cols - len(r)))
    return rows


def _build_docx(pages: list[dict[str, Any]]) -> bytes:
    if Document is None:
        raise HTTPException(status_code=500, detail="python-docx is not installed")

    doc = Document()
    for idx, p in enumerate(pages):
        doc.add_heading(f"Page {p['page'] + 1}", level=1)
        if p.get("ocr_text"):
            doc.add_paragraph(p["ocr_text"])

        md = p.get("tsr_markdown") or ""
        rows = _parse_markdown_table(md)
        if rows:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            for r_i, row in enumerate(rows):
                for c_i, cell in enumerate(row):
                    table.cell(r_i, c_i).text = cell

        if idx < len(pages) - 1:
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


class UrlPayload(BaseModel):
    url: HttpUrl


@app.post("/extract")
async def extract(
    file: UploadFile = File(...),
    task: str = Query("ocr", description="ocr|layout|tsr|all"),
    output: str = Query("json", description="json|minimal|text|txt|docx"),
    threshold: float | None = Query(None),
    x_api_key: str | None = Header(default=None, alias="X-API-Key"),
) -> dict[str, Any]:
    _require_api_key(x_api_key)

    raw = await file.read()
    _check_upload_size(raw)

    t0 = time.time()
    threshold_value = SETTINGS.default_threshold if threshold is None else float(threshold)

    if task not in {"ocr", "layout", "tsr", "all"}:
        raise HTTPException(status_code=400, detail="Invalid task. Use: ocr|layout|tsr|all")
    if output not in {"json", "minimal", "text", "txt", "docx"}:
        raise HTTPException(status_code=400, detail="Invalid output. Use: json|minimal|text|txt|docx")

    async with JOB_SEM:
        def _run() -> list[dict[str, Any]]:
            images = _load_images(file.filename or "upload", raw)
            ocr = _get_ocr() if task in {"ocr", "tsr", "all"} else None
            layout = _get_layout() if task in {"layout", "all"} else None
            tsr = _get_tsr() if task in {"tsr", "all"} else None

            pages: list[dict[str, Any]] = []
            for page_idx, img in enumerate(images):
                page: dict[str, Any] = {"page": page_idx}

                ocr_out = None
                if ocr is not None:
                    ocr_out = _ocr_page(ocr, img)
                    page["ocr_text"] = ocr_out["text"]
                    if output == "json":
                        page["ocr_lines"] = ocr_out["lines"]

                if layout is not None:
                    lyt = layout.forward([img], thr=threshold_value)[0]
                    if output == "json":
                        page["layout"] = lyt

                if tsr is not None:
                    comps = tsr([img], thr=threshold_value)[0]
                    # Rebuild OCR boxes in the structure expected by TableStructureRecognizer.construct_table.
                    ocr_boxes = []
                    if ocr_out is not None:
                        for line in ocr_out["lines"]:
                            x0, y0, x1, y1 = line["bbox"]
                            ocr_boxes.append(
                                {
                                    "x0": x0,
                                    "x1": x1,
                                    "top": y0,
                                    "bottom": y1,
                                    "text": line["text"],
                                    "layout_type": "table",
                                    "page_number": page_idx,
                                }
                            )
                    page["tsr_markdown"] = _table_markdown(img, comps, ocr_boxes)
                    if output == "json":
                        page["tsr_components"] = comps

                pages.append(page)
            return pages

        pages = await asyncio.to_thread(_run)

    dt_ms = int((time.time() - t0) * 1000)
    if output == "text":
        # Return only OCR text (best for feeding into another system/LLM).
        chunks: list[str] = []
        for p in pages:
            if "ocr_text" in p and p["ocr_text"]:
                chunks.append(str(p["ocr_text"]).strip())
        return PlainTextResponse("\n\n".join(chunks), media_type="text/plain; charset=utf-8")
    if output == "txt":
        chunks: list[str] = []
        for p in pages:
            if "ocr_text" in p and p["ocr_text"]:
                chunks.append(str(p["ocr_text"]).strip())
            if "tsr_markdown" in p and p["tsr_markdown"]:
                chunks.append(str(p["tsr_markdown"]).strip())
        content = "\n\n".join(chunks).encode("utf-8")
        return StreamingResponse(
            io.BytesIO(content),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": "attachment; filename=ocr.txt"},
        )
    if output == "docx":
        content = _build_docx(pages)
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=ocr.docx"},
        )

    if output == "minimal":
        minimal_pages: list[dict[str, Any]] = []
        for p in pages:
            mp: dict[str, Any] = {"page": p["page"]}
            if "ocr_text" in p:
                mp["ocr_text"] = p["ocr_text"]
            if "tsr_markdown" in p:
                mp["tsr_markdown"] = p["tsr_markdown"]
            minimal_pages.append(mp)
        return {"meta": {"duration_ms": dt_ms, "pages": len(minimal_pages), "task": task}, "pages": minimal_pages}

    return {"meta": {"duration_ms": dt_ms, "pages": len(pages), "task": task}, "pages": pages}


@app.post("/extract_url")
async def extract_url(
    payload: UrlPayload,
    task: str = Query("ocr", description="ocr|layout|tsr|all"),
    output: str = Query("json", description="json|minimal|text|txt|docx"),
    threshold: float | None = Query(None),
    x_api_key: str | None = Header(default=None, alias="X-API-Key"),
) -> Any:
    _require_api_key(x_api_key)

    filename, raw = _download_url(str(payload.url))
    _check_upload_size(raw)

    t0 = time.time()
    threshold_value = SETTINGS.default_threshold if threshold is None else float(threshold)

    if task not in {"ocr", "layout", "tsr", "all"}:
        raise HTTPException(status_code=400, detail="Invalid task. Use: ocr|layout|tsr|all")
    if output not in {"json", "minimal", "text", "txt", "docx"}:
        raise HTTPException(status_code=400, detail="Invalid output. Use: json|minimal|text|txt|docx")

    async with JOB_SEM:
        def _run() -> list[dict[str, Any]]:
            images = _load_images(filename, raw)
            ocr = _get_ocr() if task in {"ocr", "tsr", "all"} else None
            layout = _get_layout() if task in {"layout", "all"} else None
            tsr = _get_tsr() if task in {"tsr", "all"} else None

            pages: list[dict[str, Any]] = []
            for page_idx, img in enumerate(images):
                page: dict[str, Any] = {"page": page_idx}

                ocr_out = None
                if ocr is not None:
                    ocr_out = _ocr_page(ocr, img)
                    page["ocr_text"] = ocr_out["text"]
                    if output == "json":
                        page["ocr_lines"] = ocr_out["lines"]

                if layout is not None:
                    lyt = layout.forward([img], thr=threshold_value)[0]
                    if output == "json":
                        page["layout"] = lyt

                if tsr is not None:
                    comps = tsr([img], thr=threshold_value)[0]
                    ocr_boxes = []
                    if ocr_out is not None:
                        for line in ocr_out["lines"]:
                            x0, y0, x1, y1 = line["bbox"]
                            ocr_boxes.append(
                                {
                                    "x0": x0,
                                    "x1": x1,
                                    "top": y0,
                                    "bottom": y1,
                                    "text": line["text"],
                                    "layout_type": "table",
                                    "page_number": page_idx,
                                }
                            )
                    page["tsr_markdown"] = _table_markdown(img, comps, ocr_boxes)
                    if output == "json":
                        page["tsr_components"] = comps

                pages.append(page)
            return pages

        pages = await asyncio.to_thread(_run)

    dt_ms = int((time.time() - t0) * 1000)
    if output == "text":
        chunks: list[str] = []
        for p in pages:
            if "ocr_text" in p and p["ocr_text"]:
                chunks.append(str(p["ocr_text"]).strip())
        return PlainTextResponse("\n\n".join(chunks), media_type="text/plain; charset=utf-8")
    if output == "txt":
        chunks: list[str] = []
        for p in pages:
            if "ocr_text" in p and p["ocr_text"]:
                chunks.append(str(p["ocr_text"]).strip())
            if "tsr_markdown" in p and p["tsr_markdown"]:
                chunks.append(str(p["tsr_markdown"]).strip())
        content = "\n\n".join(chunks).encode("utf-8")
        return StreamingResponse(
            io.BytesIO(content),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": "attachment; filename=ocr.txt"},
        )
    if output == "docx":
        content = _build_docx(pages)
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=ocr.docx"},
        )

    if output == "minimal":
        minimal_pages: list[dict[str, Any]] = []
        for p in pages:
            mp: dict[str, Any] = {"page": p["page"]}
            if "ocr_text" in p:
                mp["ocr_text"] = p["ocr_text"]
            if "tsr_markdown" in p:
                mp["tsr_markdown"] = p["tsr_markdown"]
            minimal_pages.append(mp)
        return {"meta": {"duration_ms": dt_ms, "pages": len(minimal_pages), "task": task}, "pages": minimal_pages}

    return {"meta": {"duration_ms": dt_ms, "pages": len(pages), "task": task}, "pages": pages}
